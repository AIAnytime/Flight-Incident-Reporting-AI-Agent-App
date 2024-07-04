import streamlit as st
import os
import base64
from pydub import AudioSegment
from openai import OpenAI
from dotenv import load_dotenv
from crewai import Agent, Task, Crew, Process
from langchain_openai import AzureChatOpenAI
from docx import Document
from io import BytesIO
import base64

load_dotenv()

def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Flight Incident Report</a>'

# Initialize OpenAI client
groq = OpenAI(
    api_key=os.environ["GROQ_API_KEY"],
    base_url="https://api.groq.com/openai/v1"
)

llm = AzureChatOpenAI(
    openai_api_version=os.environ["OPENAI_API_GPT_4_VERSION"],
    azure_deployment="gpt-4o",
    model="gpt-4o",
    temperature=0.7,
    openai_api_key=os.environ["OPENAI_API_GPT_4_KEY"],
    azure_endpoint=os.environ["OPENAI_API_GPT_4_BASE"]
)

# Function to convert audio file to base64
def audio_to_base64(file):
    audio_file = open(file, "rb")
    audio_bytes = audio_file.read()
    base64_audio = base64.b64encode(audio_bytes).decode()
    return base64_audio

st.set_page_config(
    layout="wide",
    page_title="Flight Incident Investigation Report App"
)

st.title("Flight Incident Investigation Report App")

uploaded_file = st.file_uploader("Upload an MP3 file", type=["mp3"])

col1, col2 = st.columns(2)

with col1:
    if uploaded_file is not None:
        # Save the uploaded file to disk
        with open("uploaded_file.mp3", "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Convert the uploaded MP3 file to base64 for embedding in HTML
        base64_audio = audio_to_base64("uploaded_file.mp3")
        
        # Embed the audio file in HTML
        audio_html = f"""
            <audio controls>
                <source src="data:audio/mp3;base64,{base64_audio}" type="audio/mp3">
                Your browser does not support the audio element.
            </audio>
        """
        st.subheader("Your Uploaded Audio File")
        st.markdown(audio_html, unsafe_allow_html=True)

        if st.button("Analyze"):

            # Transcribe the audio using OpenAI API
            with open("uploaded_file.mp3", "rb") as audio_file:
                transcript = groq.audio.transcriptions.create(
                    model="whisper-large-v3",
                    file=audio_file,
                    response_format="text"
                )

            st.success("Raw Transcription: " + transcript)

            with col2:
                # Placeholder for future use or additional functionalities
                with st.spinner("Processing"):
                    # Define Agents
                    transcription_reader = Agent(
                        role="ATC Blackbox Transcription Reader",
                        goal="Read the transcription from an ATC Blackbox and clean the transcription and create a conversation flow.",
                        backstory="This agent specializes in reading the transcript from an ATC Blackbox and clean the transcript to create a conversation flow so it is easier to interpret.",
                        verbose=True,
                        allow_delegation=False,
                        llm=llm
                    )

                    incident_report_writer = Agent(
                        role="Incident Reporter",
                        goal="Generate a detailed incident investigation report based on the transcription data.",
                        backstory="This agent specializes in creating comprehensive incident investigation reports for aviation incidents. It considers the cleaned transcription data, communication logs, and aviation protocols to provide a detailed report.",
                        verbose=True,
                        allow_delegation=False,
                        llm=llm
                    )

                    # Task for the transcription reader agent
                    clean_transcription_task = Task(
                        description='Read the transcription from an ATC Blackbox, clean it, and create a conversation flow.',
                        agent=transcription_reader,
                        expected_output='Formatted transcription as conversation.',
                        inputs={'transcription': transcript}
                    )

                    # Task for the incident report writer agent
                    incident_report_task = Task(
                        description='Generate a detailed incident investigation report based on the cleaned transcription.',
                        agent=incident_report_writer,
                        expected_output='Detailed incident investigation report.'
                    )
                    # Create Crew
                    crew = Crew(
                        agents=[transcription_reader, incident_report_writer],
                        tasks=[clean_transcription_task, incident_report_task],
                        process=Process.sequential,
                        verbose=2
                    )

                    result = crew.kickoff()
                    st.subheader("Incident Report")
                    st.write(result)


